import subprocess, sys
try:
    from docx import Document
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)

for h_name, size, bold in [("Heading 1", 16, True), ("Heading 2", 13, True), ("Heading 3", 11, True)]:
    s = styles[h_name]
    s.font.name  = "Arial"
    s.font.size  = Pt(size)
    s.font.bold  = bold
    s.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# ── Helper functions ──────────────────────────────────────────────────────────
def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_para(text, bold_prefix=None, space_after=6):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        p.add_run(" " + text).font.name = "Arial"
    else:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        p.add_run(" " + text).font.name = "Arial"
    else:
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p

def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

def make_table(headers, rows, col_widths, header_fill="1F497D", header_text_color="FFFFFF"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hdr = t.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr.cells[i]
        cell.width = Inches(w)
        shade_cell(cell, header_fill)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(*bytes.fromhex(header_text_color))
        run.font.name = "Arial"
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(0)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        fill = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, (val, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[ci]
            cell.width = Inches(w)
            shade_cell(cell, fill)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()  # spacing after table

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("CLV Prediction — Technical Development Log")
run.bold = True
run.font.name  = "Arial"
run.font.size  = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
title_p.paragraph_format.space_after = Pt(6)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("Advanced Analytics Assignment 1  |  Group: Lowie")
run.font.name  = "Arial"
run.font.size  = Pt(13)
run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
sub_p.paragraph_format.space_after = Pt(20)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. Project Overview", 1)
add_para("Goal: Predict 2018–2019 revenue (Customer Lifetime Value) for customers of a "
         "European fashion shoe shop, using only their 2016–2017 transaction history.")

add_para("Competition metrics:", space_after=3)
add_bullet("Primary: Mean Absolute Error (MAE) — lower is better")
add_bullet("Secondary: Spearman rank correlation — higher is better")
add_bullet("Professor benchmark: MAE 61.146 / Spearman 0.41")

add_para("Data:", space_after=3)
add_bullet("Training set: ~116,000 customers with known 2018–2019 revenue labels")
add_bullet("Test set: ~29,000 customers (no labels)")
add_bullet("Transaction data: ~1.2 M rows covering 2016–2017")
add_bullet("Key challenge: 63.4% of customers churned (revenue = 0 in 2018–2019), "
           "making this a highly zero-inflated regression problem")

# ══════════════════════════════════════════════════════════════════════════════
# 2. APPROACH EVOLUTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2. Approach Evolution", 1)

# 2.1
add_heading("2.1  Initial Two-Stage Pipeline", 2)
add_para("Architecture: Binary churn classifier → revenue regressor.")
add_para("Stage 1 — Churn classifier:", bold_prefix="Stage 1 — Churn classifier:",
         space_after=3)
add_para("Three gradient boosting models (LightGBM, XGBoost, CatBoost) trained to predict "
         "whether a customer would return (revenue > 0). Ensemble averaged probabilities.")
add_para("Stage 2 — Revenue regressor:", bold_prefix="Stage 2 — Revenue regressor:",
         space_after=3)
add_para("Trained on returning customers only. Predicted log1p(revenue) with MAE objective.")
add_para("Prediction: If P(return) > threshold → predict revenue × P(return), else predict 0. "
         "Threshold grid-searched over [0.10, 0.90].")
add_para("Results: MAE 62.65 / Spearman 0.3964", bold_prefix="Results:")

add_para("Problems identified:", space_after=3)
add_bullet("Churn classifier AUC 0.72 — churner/returner distributions overlap heavily in RFM feature space.")
add_bullet("Hard threshold creates 'arms' in the scatter plot: horizontal arm (true returners predicted as churned) "
           "and vertical arm (true churners predicted as returning). Both inflate MAE directly.")
add_bullet("No feature set can fully resolve this ambiguity — churn in fashion retail is inherently probabilistic.")

# 2.2
add_heading("2.2  Probability Calibration", 2)
add_para("Change: Replaced deprecated CalibratedClassifierCV(cv='prefit') with IsotonicRegression "
         "fitted on a 20% calibration split (three-way split: 60% train / 20% calibration / 20% validation).")
add_para("Why it matters: Raw classifier probabilities are overconfident. Isotonic regression maps "
         "them to calibrated posteriors without changing AUC (calibration is a monotonic transformation).")
add_para("Results: Log loss improved 0.6013 → 0.5758. AUC unchanged at 0.72. "
         "MAE 62.50 / Spearman 0.4002. Small improvement; arm problem remained.", bold_prefix="Results:")

# 2.3
add_heading("2.3  Tweedie Regression  (Discarded)", 2)
add_para("Idea: Replace the two-stage architecture with a single Tweedie distribution model, "
         "which naturally handles zero-inflated positive data.")
add_para("Why it failed: Tweedie regression never predicts exactly 0. The 63.4% of churned customers "
         "all receive small positive predictions, systematically inflating MAE.")
add_para("Results: MAE 69.88 — significantly worse. Discarded entirely.", bold_prefix="Results:")

# 2.4
add_heading("2.4  Feature Engineering", 2)
add_para("Original: Basic RFM aggregates (frequency, recency, monetary value, return rates). "
         "The following were added iteratively:")

add_bullet("Quarterly breakdown (18 features): 6 quarters (2016Q1–2017Q2) × 3 metrics "
           "(revenue, orders, returns). Lets the model weight recent quarters more heavily — "
           "a customer active in 2017Q2 is a stronger retention signal than one active only in 2016Q1.",
           bold_prefix="Quarterly breakdown (18 features):")
add_bullet("Recency normalisation: recency_normalized = recency_days / (avg_days_between_orders + 1), "
           "clipped at 99th percentile. A 60-day gap means different things for a weekly vs annual shopper.",
           bold_prefix="Recency normalisation:")
add_bullet("Target encoding: brand_target_enc and category_target_enc via 5-fold CV with Laplace "
           "smoothing (MIN_SAMPLES=5) to prevent label leakage.",
           bold_prefix="Target encoding:")
add_bullet("Return rates: brand_return_rate and category_return_rate from transaction data.",
           bold_prefix="Return rates:")

add_para("Total features after engineering: 74. Impact was marginal on final MAE — the bottleneck "
         "was churn identification, not revenue prediction for identified returners.")

# 2.5
add_heading("2.5  BG/NBD Model", 2)
add_para("Model: Beta-Geometric / Negative Binomial Distribution (Fader, Hardie & Lee, 2005) — "
         "the academic standard for CLV in non-contractual settings.")
add_para("Two components:", space_after=3)
add_bullet("BG/NBD: Models purchase frequency (Poisson) and permanent dropout (geometric). "
           "Outputs p_alive — the posterior probability a customer has NOT permanently dropped out.",
           bold_prefix="BG/NBD:")
add_bullet("Gamma-Gamma: Models expected revenue per transaction conditional on activity. "
           "Fitted only on customers with ≥ 1 repeat purchase.",
           bold_prefix="Gamma-Gamma:")

add_para("Implementation: lifetimes library (BetaGeoFitter + GammaGammaFitter). "
         "Time unit: weeks (days cause numerical instability). penalizer_coef=0.001.")
add_para("Technical fix: lifetimes stores a local lambda (generate_new_data) that pickle cannot "
         "serialise. Fixed by popping it before joblib.dump and restoring it after.")
add_para("8 features extracted: p_alive, exp_purchases_24m, exp_avg_revenue, bgnbd_clv, "
         "bgnbd_frequency, bgnbd_recency_weeks, bgnbd_T_weeks, bgnbd_monetary_value. "
         "Merged with 74 engineered features → 82 features total (customer_features_v3.csv).")

# 2.6
add_heading("2.6  Single-Stage ML with BG/NBD Features", 2)
add_para("Architecture change: Single model trained on ALL customers (including churners with revenue = 0).")
add_para("Target: log1p(revenue). Loss: MAE objective (regression_l1 / reg:absoluteerror / MAE). "
         "With MAE loss, each tree leaf predicts the median of its samples — leaves dominated "
         "by churners naturally predict 0 without needing a separate classifier or threshold.")
add_para("Methodology:", space_after=3)
add_bullet("40-iteration RandomizedSearchCV + 5-fold CV per model for hyperparameter search")
add_bullet("5-fold out-of-fold (OOF) predictions for robust MAE estimation without leakage")
add_bullet("Nelder-Mead optimisation with softmax parameterisation for ensemble weights (LGB/XGB/CAT)")
add_bullet("Scalar alpha blend: final = alpha × bgnbd_clv + (1-alpha) × ML_ensemble, alpha optimised on OOF")
add_para("Results: MAE ~62.50 / Spearman ~0.40 — similar to two-stage pipeline.", bold_prefix="Results:")
add_para("Why similar: The bottleneck is feature-space overlap, not architecture. The arms persisted "
         "in the scatter plot, confirming the model faces the same fundamental ambiguity.")

# 2.7
add_heading("2.7  P_alive Soft Scaling  (Discarded)", 2)
add_para("Idea: prediction = p_alive × E[revenue | customer is active]. Train revenue model on "
         "returning customers only, predict for all, multiply by p_alive — mirrors BG/NBD + "
         "Gamma-Gamma but with a stronger ML revenue estimator.")
add_para("Results: OOF MAE 109.46 — significantly worse than ensemble. Discarded.", bold_prefix="Results:")
add_para("Why it failed: p_alive is already one of the 82 input features. The trees learned the "
         "optimal relationship between p_alive and the target implicitly. Multiplying by it again "
         "adds distortion. The conditional model (trained only on returners) also cannot learn "
         "when to predict near-zero.")

# 2.8
add_heading("2.8  Optuna Hyperparameter Optimisation", 2)
add_para("Motivation: RandomizedSearchCV samples blindly. Optuna's TPE (Tree-structured Parzen "
         "Estimator) sampler learns which hyperparameter regions produce low loss and focuses "
         "exploration there — far more efficient with 8+ interacting parameters.")
add_para("Setup:", space_after=3)
add_bullet("60 trials for LightGBM and XGBoost, 40 for CatBoost (slower per trial)")
add_bullet("3-fold CV inside Optuna (fast proxy); full 5-fold OOF with winning params")
add_bullet("Wider search ranges: learning_rate 0.005–0.15 (log), n_estimators 300–3000, "
           "num_leaves 20–300, reg_alpha/lambda 1e-8 to 10 (log), XGBoost adds gamma parameter")
add_para("Best params found:", space_after=3)
add_bullet("LightGBM: n_estimators=1441, lr=0.00877, num_leaves=35, min_child_samples=77")
add_bullet("XGBoost: n_estimators=466, lr=0.01059, max_depth=6, min_child_weight=20")
add_bullet("CatBoost: iterations=1761, lr=0.00820, depth=8")
add_bullet("Ensemble weights (Nelder-Mead optimised): LGB=0.202, XGB=0.256, CAT=0.542")
add_para("Results: OOF MAE 62.42 / Val MAE 62.37 / Val Spearman 0.3706. All three algorithms "
         "converged to essentially the same OOF score (LGB 62.56, XGB 62.57, CAT 62.46). "
         "Three independent algorithms returning the same MAE despite very different hyperparameters "
         "is a diagnostic signal: the bottleneck is not hyperparameters, it is the information "
         "available in the features.", bold_prefix="Results:")
add_para("BG/NBD blend: Optimal alpha = 0.0 — BG/NBD adds nothing on top of the ML ensemble. "
         "The 8 BG/NBD features are already inside the 82-feature input space, so the trees "
         "have already learned the BG/NBD signal implicitly.")

# 2.9
add_heading("2.9  Kevin's EDA Features  (Teammate Contribution)", 2)
add_para("Source: Teammate Kevin (colin0019) pushed 4 R EDA notebooks analysing the transaction "
         "data from different angles. Three groups of features not yet in the Python pipeline "
         "were identified and integrated into 02_feature_engineering.ipynb.")
add_para("9 new features added across 3 groups:", space_after=3)
add_bullet("Sales timing (3 features): pct_orders_during_sales (fraction of orders placed in "
           "January or July — ~25% of customers only shop during sale months), revenue_trend "
           "(revenue_2017/revenue_2016 ratio; −999 sentinel if no 2016 baseline), "
           "frequency_trend (orders_2017/orders_2016 ratio).",
           bold_prefix="Sales timing (3 features):")
add_bullet("Gender/product segment (3 features): women_items_only, men_items_only, "
           "children_items_only — flags whether all items a customer bought belonged to a single "
           "gender segment. Distinguishes personal shoppers from gift buyers.",
           bold_prefix="Gender/product segment (3 features):")
add_bullet("Return shop behaviour (3 features): returned_to_flagship (ever returned to the "
           "main flagship store ID), returned_to_high_loyalty_shop (returned to any of 6 shops "
           "with above-average repeat purchase rates), n_return_shops (number of distinct return "
           "shop IDs used).",
           bold_prefix="Return shop behaviour (3 features):")
add_para("Total features: 74 + 9 new = 83 in v2; 83 + 8 BG/NBD = 90 in v3 "
         "(customer_features_v3.csv).")
add_para("Results: Val MAE 62.37 / Val Spearman 0.3706 — marginal MAE improvement (~0.08) "
         "over the pre-Kevin ensemble (62.45). Feature ceiling confirmed: more features "
         "continue to yield diminishing returns.", bold_prefix="Results:")

# 2.10
add_heading("2.10  Two-Stage ML: Churn Classifier × Conditional Revenue", 2)
add_para("Root cause: 63.4% of training customers have exactly €0 revenue yet all single-stage "
         "models predict a small positive number for nearly all of them. In log1p space the "
         "churner target (log1p(0) = 0) and a low-revenue returner target sit very close "
         "together — the decision boundary is inherently blurry. All three Optuna-tuned "
         "algorithms converging to OOF MAE ≈ 62.42 confirmed this is a feature-space ceiling, "
         "not a hyperparameter problem.")
add_para("Architecture:", space_after=3)
add_bullet("Stage 1 — Churn classifier: LightGBM binary classifier trained on all customers, "
           "all 90 features. Target: is_returner = (revenue_2018_2019 > 0). "
           "Output: P(customer returns in 2018–2019).",
           bold_prefix="Stage 1 — Churn classifier:")
add_bullet("Stage 2 — Conditional regressor: Optuna-tuned LightGBM regressor trained on "
           "returning customers only. Target: log1p(revenue). "
           "Output: E[revenue | customer is a returner].",
           bold_prefix="Stage 2 — Conditional regressor:")
add_bullet("Combination: prediction = P(return) × conditional_revenue.",
           bold_prefix="Combination:")
add_bullet("Optional threshold: an OOF-optimised revenue threshold (77.1€) that zeros out "
           "residual predictions below the break-even level.",
           bold_prefix="Optional threshold:")

add_para("Comparison to original two-stage (section 2.1):", space_after=4)
make_table(
    headers=["Dimension", "Original two-stage (2.1)", "New two-stage (2.10)"],
    col_widths=[1.8, 2.7, 2.5],
    rows=[
        ["Classifier ensemble",   "LGB + XGB + CatBoost (3 models)",       "LGB only"],
        ["Tuning metric",         "AUC — discriminative, not MAE-aligned",  "Binary crossentropy + MAE loss"],
        ["Prob. calibration",     "Isotonic regression on 20% cal set",     "None (raw probabilities)"],
        ["Decision rule",         "Hard: P > threshold → predict, else 0", "Soft: P × conditional_revenue"],
        ["Feature set",           "~51 features",                          "90 features (v3)"],
        ["Stage integration",     "Separate notebooks, separate splits",   "Both stages inside same 5-fold OOF"],
        ["Evaluation",            "AUC + MAE separately",                  "Single OOF MAE for combined product"],
    ]
)
add_para("OOF results:", space_after=3)
add_bullet("Two-stage ML: OOF MAE 71.49 / Spearman 0.3985 — better Spearman than ensemble (0.3784) "
           "but much worse MAE (71.49 vs 62.42).")
add_bullet("Two-stage + threshold (77.1€): OOF MAE 62.82 / Spearman 0.3937 — threshold helps "
           "but still worse than pure ensemble on MAE.")
add_para("Why the trade-off: P(return) < 1 for every customer, so multiplying conditional_revenue "
         "by a probability always scales predictions down. This compresses the revenue magnitude "
         "for true returners, hurting MAE, but the classifier correctly re-orders customers "
         "(churners rank lower, high-value returners rank higher) — hence better Spearman.")
add_para("Decision: USE_TWOSTAGE=False. Pure ML ensemble retained as the MAE-optimal base.")

# 2.11
add_heading("2.11  Two-Stage × ML Ensemble Blend  (Final Approach)", 2)
add_para("Insight: The two-stage has better Spearman (ranks customers more correctly) while the "
         "ML ensemble has better MAE (better absolute magnitude). A blend of the two can "
         "harvest both strengths: final = β × two_stage + (1−β) × ml_ensemble.")
add_para("Optimisation: Sweep β ∈ [0, 1] on OOF predictions. The MAE curve is strictly "
         "monotone increasing in β — the optimizer finds β*=0 (pure ML ensemble minimises MAE). "
         "However, the Spearman curve reveals a sharp cliff: at β≈0.04 Spearman jumps from "
         "0.3784 to ~0.401 on OOF, then stays roughly flat up to β≈0.6. Even 4% two-stage "
         "weight dramatically re-orders customers near the churn boundary — where the Spearman "
         "ranking matters most — at negligible MAE cost.")
add_para("Decision: Manually fix β=0.04. The MAE-optimal β (0) ignores Spearman entirely; "
         "at β=0.04 we gain ~+0.023 Spearman for only +0.03 OOF MAE.", bold_prefix="Decision:")
add_para("Val results (held-out 20%):", space_after=3)
add_bullet("Two-stage × ML blend (β=0.04): Val MAE 62.69 / Val Spearman 0.4032")
add_bullet("vs. pure ML ensemble: Val MAE 62.37 / Val Spearman 0.3706")
add_bullet("Cost: +0.32 MAE. Gain: +0.033 Spearman (Spearman now within 0.007 of benchmark 0.41)")
add_para("This is the final submission strategy. USE_TS_ML_BLEND=True saved to models/blend_beta.pkl "
         "and picked up by 05_test_prediction.ipynb → submission_v7_ts_ml_blend_b0.040.csv.")

# ══════════════════════════════════════════════════════════════════════════════
# 3. RESULTS SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3. Results Summary", 1)
add_para("All MAE/Spearman figures are on the held-out 20% validation set unless noted as OOF.", space_after=4)
make_table(
    headers=["Approach", "Val MAE", "Spearman", "Notes"],
    col_widths=[2.9, 0.85, 0.9, 2.35],
    rows=[
        ["Two-stage: ML churn + revenue",            "62.65",  "0.3964", "Baseline (section 2.1)"],
        ["+ Probability calibration",                "62.50",  "0.4002", "IsotonicRegression (2.2)"],
        ["Tweedie regression",                       "69.88",  "—",      "Discarded: never predicts 0 (2.3)"],
        ["Single-stage + 82 features (BG/NBD)",      "~62.50", "~0.40",  "BG/NBD features added (2.5–2.6)"],
        ["P_alive soft scaling",                     "worse",  "—",      "Discarded: p_alive already a feature (2.7)"],
        ["Optuna HPO — ML ensemble (90 features)",   "62.37",  "0.3706", "OOF MAE 62.42; all 3 algorithms converged (2.8)"],
        ["+ Kevin's EDA features (9 new)",           "62.37",  "0.3706", "−0.08 MAE; feature ceiling confirmed (2.9)"],
        ["Two-stage ML (clf × regressor)",           "71.55",  "0.4009", "Better Spearman; worse MAE (2.10)"],
        ["Two-stage + threshold (77€)",              "62.92",  "0.3937", "Discarded: ML ensemble still better (2.10)"],
        ["Two-stage × ML blend β=0.04  ← FINAL",    "62.69",  "0.4032", "0.007 below benchmark Spearman (2.11)"],
        ["Professor benchmark",                      "61.146", "0.41",   "Target"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. TECHNICAL CHALLENGES
# ══════════════════════════════════════════════════════════════════════════════
add_heading("4. Technical Challenges & Fixes", 1)
make_table(
    headers=["Problem", "Cause", "Fix"],
    col_widths=[2.2, 2.2, 2.6],
    rows=[
        ["CalibratedClassifierCV(cv='prefit') error",
         "Removed in newer scikit-learn",
         "IsotonicRegression fitted directly on calibration set"],
        ["lifetimes PicklingError on joblib.dump",
         "Local lambda (generate_new_data) in fitted model",
         "Pop attribute before dump, restore after"],
        ["SHAP shape mismatch with CatBoost",
         "CatBoost adds a bias column to SHAP output",
         "shap_values[:, :-1] to strip bias column"],
        ["X_val_ret undefined in notebook 06",
         "Variable renamed during refactor",
         "Replaced with X_shap throughout cells 7–9"],
        ["OSError: Too many levels of symbolic links",
         "data/, models/, submissions/ were git-tracked\nself-referential symlinks",
         "rm symlinks; mkdir -p data models submissions;\nre-add data files directly"],
        ["NameError: 'idx' not defined in notebook 06",
         "f-strings inside a list literal are evaluated\nwhen the list is built, before the for loop runs",
         "Moved the f-string title inside the loop body\nso idx is already assigned"],
        ["Notebook 06 reporting stale MAE/Spearman",
         "Error-analysis cell loaded USE_BLEND but not\nUSE_TS_ML_BLEND or blend_beta",
         "Updated setup + error_residuals cells to load\nboth flags and compute val_pred_ts_blend"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. ARCHITECTURE & DESIGN DECISIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("5. Architecture & Design Decisions", 1)

add_para("Why log1p(revenue) as target?", bold_prefix="Why log1p(revenue) as target?")
add_para("Revenue is zero-inflated and right-skewed. log1p compresses the scale (€1000→6.9, "
         "€100→4.6, €0→0), reducing outlier influence and helping trees split evenly. "
         "The 0→0 mapping preserves the natural zero boundary.")

add_para("Why MAE loss?", bold_prefix="Why MAE loss?")
add_para("The competition evaluates on MAE. MAE loss (predicting the leaf median) directly aligns "
         "training and evaluation. RMSE would optimise the mean, overfitting to high-revenue outliers.")

add_para("Why weekly time units for BG/NBD?", bold_prefix="Why weekly time units for BG/NBD?")
add_para("BG/NBD assumes a Poisson purchase process. With daily units and customers buying "
         "only monthly/quarterly, parameter estimation becomes numerically unstable. "
         "Weeks strike the right balance.")

add_para("Why 5-fold OOF for ensemble weights?", bold_prefix="Why 5-fold OOF for ensemble weights?")
add_para("Using a held-out validation set to find weights risks overfitting that specific split. "
         "OOF covers all training customers and gives a less biased MAE estimate.")

add_para("Why Nelder-Mead with softmax?", bold_prefix="Why Nelder-Mead with softmax?")
add_para("Constrained optimisation (weights sum to 1, all positive) becomes unconstrained "
         "via softmax over logits. Nelder-Mead is derivative-free and robust to the non-smooth MAE objective.")

add_para("Why manually fix β=0.04 instead of the MAE-optimal β=0?",
         bold_prefix="Why manually fix β=0.04 instead of the MAE-optimal β=0?")
add_para("The competition has two metrics: MAE and Spearman. The optimizer finds β=0 because it "
         "minimises only MAE. The Spearman sweep reveals a sharp non-linear cliff: even 4% "
         "two-stage weight re-orders borderline customers (where Spearman sensitivity is highest) "
         "at almost no MAE cost. A β that Pareto-dominates on the combined objective space is "
         "the right choice when both metrics are evaluated.")

# ══════════════════════════════════════════════════════════════════════════════
# 6. NOTEBOOK STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("6. Notebook Structure", 1)
make_table(
    headers=["Notebook", "Purpose"],
    col_widths=[2.8, 4.2],
    rows=[
        ["02_feature_engineering.ipynb",
         "RFM + quarterly + recency_normalized + target encoding + Kevin's 9 EDA features "
         "→ customer_features_v2.csv (83 features)"],
        ["03_bgnbd_model.ipynb",
         "BG/NBD fit, p_alive extraction, merge → customer_features_v3.csv (90 features) + bgf/ggf models"],
        ["03_churn_model.ipynb",
         "Old ML churn classifier (kept for reference, not in active pipeline)"],
        ["04_revenue_model.ipynb",
         "Optuna HPO, 5-fold OOF, Nelder-Mead ensemble, soft scaling comparison, "
         "two-stage ML OOF, β sweep blend optimisation, retrain all models, save"],
        ["05_test_prediction.ipynb",
         "Load best models, apply β=0.04 blend strategy, predict test set "
         "→ submission_v7_ts_ml_blend_b0.040.csv"],
        ["06_interpretability.ipynb",
         "SHAP beeswarm, p_alive dependence plot, quarterly importance, "
         "waterfall plots, error analysis (applies blend strategy)"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. KEY LESSONS LEARNED
# ══════════════════════════════════════════════════════════════════════════════
add_heading("7. Key Lessons Learned", 1)

lessons = [
    ("Architecture bottleneck vs feature bottleneck",
     "The arms problem is not caused by two-stage vs single-stage — it is caused by "
     "inherent overlap between churner and returner feature distributions. No architecture "
     "can cleanly separate what the data cannot."),
    ("Calibration ≠ discrimination",
     "Isotonic calibration improved log loss but not AUC or MAE, because it only "
     "rescales probabilities monotonically."),
    ("Soft scaling can hurt when the signal is already embedded",
     "p_alive was already a feature. Multiplying by it again added distortion rather than signal."),
    ("BG/NBD is best as a feature, not a replacement",
     "The probabilistic model provides valuable p_alive and CLV estimates that tree models "
     "can use contextually. Using BG/NBD alone loses the 74+ engineered features. "
     "BG/NBD blend alpha optimised to 0.0, confirming the ML ensemble fully subsumes it."),
    ("MAE in log-space ≠ MAE in original space",
     "Training on log1p(revenue) with MAE loss minimises median log1p error, which approximately "
     "minimises original-space MAE. Far better than raw revenue (where leaf medians are mostly 0)."),
    ("RandomizedSearchCV is insufficient for 8+ interacting hyperparameters",
     "40 random trials cover a tiny fraction of a high-dimensional space. "
     "Optuna TPE makes each trial informative for the next."),
    ("Algorithm convergence is a ceiling signal",
     "When LightGBM, XGBoost, and CatBoost all return the same OOF MAE after independent "
     "Optuna searches, further tuning will not help — the bottleneck is the feature space, "
     "not the model. The right response is an architectural change, not more trials."),
    ("OOF integration is critical for honest two-stage evaluation",
     "The original two-stage used separate train/val splits for the classifier and regressor, "
     "so the combined P × revenue MAE was never properly measured. Only by running both stages "
     "inside the same OOF fold can the combined prediction be compared apples-to-apples "
     "against a single-stage baseline."),
    ("MAE-optimal ≠ jointly optimal when two metrics are evaluated",
     "Optimising β solely for MAE finds β=0 (pure ML ensemble). The Spearman sweep reveals "
     "a sharp non-linear cliff at β=0.04: Spearman jumps +0.023 at almost no MAE cost. "
     "When both MAE and Spearman are evaluated, the single-metric optimum is not the right choice."),
    ("OOF Spearman ≠ val Spearman — but the direction is consistent",
     "OOF Spearman for the pure ML ensemble was 0.3784; on the held-out val set it was 0.3706. "
     "The gap (~0.008) is expected — OOF covers 93k training customers while the val set covers "
     "23k held-out customers. The blend improvement (+0.033 Spearman on val) matched the "
     "OOF prediction direction, confirming the sweep chart was reliable."),
]

for i, (title, body) in enumerate(lessons, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{i}.  {title}: ")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r2 = p.add_run(body)
    r2.font.name = "Arial"
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = ("/Users/lowievanlitsenborg/Library/Mobile Documents/"
       "com~apple~CloudDocs/lessen/1e Ma/Advanced Analytics/Assignments/1/"
       "AdvancedAnalyticsTabular/Code_Lowie/Development_Report.docx")
doc.save(out)
print(f"Saved: {out}")
